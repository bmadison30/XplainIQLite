# XplainIQLite_Final.py
# Complete production-ready version with client submission and admin report generation

import io
import base64
import time
from datetime import datetime
from typing import Optional, Dict, List, Tuple
from pathlib import Path

import streamlit as st
import pandas as pd

# Optional chart (admin radar)
try:
    import matplotlib as mpl
    import matplotlib.pyplot as plt
    HAS_MPL = True
    mpl.rcParams.update({
        "font.size": 10,
        "axes.titlesize": 12,
        "axes.labelsize": 10,
        "xtick.labelsize": 9,
        "ytick.labelsize": 9,
    })
except Exception:
    HAS_MPL = False

# Optional DOCX report
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# Optional PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    HAS_REPORTLAB = True
except Exception:
    HAS_REPORTLAB = False

# =========================
# Paths & Brand
# =========================
try:
    APP_DIR = Path(__file__).resolve().parent
except NameError:
    APP_DIR = Path.cwd()

ASSETS_DIR = APP_DIR / "assets"
ASSETS_DIR.mkdir(exist_ok=True)

BRAND_NAME = "XplainIQ lite: Channel Readiness Scoring Index"
ASSET_LOGO_PATH = ASSETS_DIR / "xplainiq_logo.png"
ASSET_FAVICON_PATH = ASSETS_DIR / "xplainiq_favicon.png"

def _load_bytes(path: Path) -> Optional[bytes]:
    try:
        if path.exists() and path.is_file():
            return path.read_bytes()
    except Exception as e:
        if 'logo_debug' not in st.session_state:
            st.session_state.logo_debug = []
        st.session_state.logo_debug.append(f"Failed to load {path}: {e}")
    return None

# ====================================
# Logo rendering
# ====================================
def _detect_mime_from_bytes(b: bytes) -> Optional[str]:
    if not b or len(b) < 12:
        return None
    head16 = b[:16]
    head32 = b[:32]
    if head16.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if head16.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if head16[:4] == b"RIFF" and b"WEBP" in head16[8:12]:
        return "image/webp"
    if head16.startswith(b"GIF87a") or head16.startswith(b"GIF89a"):
        return "image/gif"
    if head16.startswith(b"BM"):
        return "image/bmp"
    if head16[:4] == b"\x00\x00\x01\x00":
        return "image/x-icon"
    if head16.startswith(b"II*\x00") or head16.startswith(b"MM\x00*"):
        return "image/tiff"
    if head32.lstrip().startswith(b"<svg") or b"<svg" in b[:200]:
        return "image/svg+xml"
    return None

def _st_html_img(bytes_data: bytes, mime: str, width: int):
    b64 = base64.b64encode(bytes_data).decode("utf-8")
    st.markdown(
        f'<img src="data:{mime};base64,{b64}" width="{width}" style="width:{width}px;height:auto;" />',
        unsafe_allow_html=True
    )

def show_logo_any(path: Path, width: int = 220, show_debug: bool = False):
    data = _load_bytes(path)
    if not data:
        if show_debug:
            st.warning(f"⚠️ Logo not found at: {path}")
            st.info(f"Current working directory: {Path.cwd()}")
            st.info(f"APP_DIR: {APP_DIR}")
            if hasattr(st.session_state, 'logo_debug'):
                for msg in st.session_state.logo_debug:
                    st.caption(msg)
        return
    
    mime = _detect_mime_from_bytes(data) or {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".webp": "image/webp", ".gif": "image/gif", ".bmp": "image/bmp",
        ".ico": "image/x-icon", ".tiff": "image/tiff", ".tif": "image/tiff",
        ".svg": "image/svg+xml",
    }.get(path.suffix.lower())
    
    if mime in ("image/png", "image/jpeg"):
        try:
            st.image(data, width=width)
            return
        except Exception:
            pass
    if mime:
        _st_html_img(data, mime, width)

def _favicon_image():
    try:
        from PIL import Image
    except Exception:
        return None

    fav_bytes = _load_bytes(ASSET_FAVICON_PATH)
    if fav_bytes:
        try:
            return Image.open(io.BytesIO(fav_bytes))
        except Exception:
            pass

    logo_bytes = _load_bytes(ASSET_LOGO_PATH)
    if not logo_bytes:
        return None
    mime = _detect_mime_from_bytes(logo_bytes)
    if mime == "image/svg+xml":
        return None
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(logo_bytes)).convert("RGBA")
        size = 32
        ratio = min(size / img.width, size / img.height)
        new_w, new_h = max(1, int(img.width * ratio)), max(1, int(img.height * ratio))
        canvas = Image.new("RGBA", (size, size), (0, 0, 0, 0))
        img = img.resize((new_w, new_h), Image.LANCZOS)
        canvas.paste(img, ((size - new_w) // 2, (size - new_h) // 2), img)
        return canvas
    except Exception:
        return None

# =====
# UI (set page config FIRST)
# =====
st.set_page_config(
    page_title="XplainIQ lite – Channel Readiness Scoring Index",
    page_icon=_favicon_image() or "🔎",
    layout="centered"
)

# =====
# App constants
# =====
PILLARS: List[Tuple[str, List[str]]] = [
    ("A. Channel Strategy & Alignment", ["A1", "A2"]),
    ("B. Partner Program Design",      ["B1", "B2"]),
    ("C. Partner Enablement & Engagement", ["C1", "C2"]),
    ("D. Sales & Operations Integration",  ["D1", "D2"]),
    ("E. Growth Readiness",            ["E1", "E2"]),
]

QUESTIONS: Dict[str, str] = {
    "A1": "Do you have a clearly defined purpose for selling through partners (beyond revenue expansion)?",
    "A2": "Are your targeted partner types (TA, VAR, MSP, SI, etc.) well-defined and prioritized?",
    "B1": "Do you have a partner program with tiering, incentives, rules of engagement, or performance criteria?",
    "B2": "Can you clearly articulate what makes your offer unique and profitable for partners?",
    "C1": "Do you provide training, sales playbooks, or co-branded marketing assets?",
    "C2": "How consistently do you communicate and collaborate with active partners?",
    "D1": "Are internal sales/ops aligned to support channel transactions (quoting, order flow, support)?",
    "D2": "Do you track partner pipeline separately with forecast accuracy goals?",
    "E1": "Does senior leadership actively sponsor the channel model?",
    "E2": "Are tools, systems, and staffing sufficient to support 2–3× partner growth?",
}

TIER_BANDS = [
    ("Emerging",   0, 39),
    ("Developing", 40, 59),
    ("Established",60, 79),
    ("Optimized",  80, 100),
]

# =====
# Scoring logic
# =====
def tier_for(score: float) -> str:
    s = round(score)
    for name, lo, hi in TIER_BANDS:
        if lo <= s <= hi:
            return name
    return "Unknown"

def pillar_commentary(pillar_name: str, pscore: float) -> str:
    if pscore >= 80:
        return f"{pillar_name} is strong and scalable – keep reinforcing what works."
    if pscore >= 60:
        return f"{pillar_name} shows a solid foundation with room to standardize and scale."
    if pscore >= 40:
        return f"{pillar_name} is emerging – formalize structure, cadence, and measurement."
    return f"{pillar_name} is underdeveloped – prioritize core mechanics and minimum viable structure."

def compute_scores(answers: Dict[str, int]) -> Tuple[List[Tuple[str, float, Dict[str, int]]], float]:
    pillar_scores = []
    for pname, qids in PILLARS:
        vals = [int(answers.get(q, 0)) for q in qids]
        if not vals or all(v == 0 for v in vals):
            pscore = 0.0
        else:
            pscore = (sum(vals) / len(vals)) / 5.0 * 100.0
        pillar_scores.append((pname, pscore, dict(zip(qids, vals))))
    overall = sum(p[1] for p in pillar_scores) / len(pillar_scores)
    return pillar_scores, overall

def derive_strengths_gaps(ps: List[Tuple[str, float, Dict[str, int]]]) -> Tuple[List[str], List[str]]:
    sorted_p = sorted(ps, key=lambda x: x[1], reverse=True)
    strengths = [p[0] for p in sorted_p[:2]]
    gaps = [p[0] for p in sorted_p[-3:]]
    return strengths, gaps

def recommend_actions(ps: List[Tuple[str, float, Dict[str, int]]]) -> List[str]:
    playbook = {
        "A. Channel Strategy & Alignment": "Clarify the partner role by segment and set a 12-month channel thesis with 3 measurable outcomes.",
        "B. Partner Program Design": "Publish a simple one-pager: tiers, incentives, rules of engagement, and co-marketing paths.",
        "C. Partner Enablement & Engagement": "Stand up a 30-60-90 enablement cadence: onboarding kit, monthly enablement call, quarterly MDF campaign.",
        "D. Sales & Operations Integration": "Separate channel pipeline tracking; define lead routing/quoting SLAs; add 'channel' to forecast reviews.",
        "E. Growth Readiness": "Baseline partner P&L and capacity; set tooling minimums (PRM/CRM views) and resource triggers for 2–3× growth."
    }
    lows = sorted(ps, key=lambda x: x[1])[:3]
    return [playbook.get(p[0], f"Prioritize foundational improvements in {p[0].lower()} to enable scale.") for p in lows]

# =====
# Radar chart
# =====
def radar_chart(pillar_scores: List[Tuple[str, float, Dict[str, int]]], title: str = "Readiness Radar"):
    import numpy as np
    from textwrap import wrap

    labels = [p[0].split(". ", 1)[1] if ". " in p[0] else p[0] for p in pillar_scores]
    values = [float(p[1]) for p in pillar_scores]

    labels_wrapped = ["\n".join(wrap(lbl, width=18)) for lbl in labels]
    values_closed = values + values[:1]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = plt.subplot(111, polar=True)
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels_wrapped, fontsize=10)
    ax.set_rlabel_position(0)
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels([20, 40, 60, 80, 100], fontsize=9)

    ax.grid(True, linestyle="--", linewidth=0.7, alpha=0.6)
    for spine in ax.spines.values():
        spine.set_alpha(0.3)

    ax.plot(angles, values_closed, linewidth=2.2)
    ax.fill(angles, values_closed, alpha=0.10)
    ax.scatter(angles, values_closed, s=25)

    for ang, val in zip(angles[:-1], values):
        ax.text(ang, min(val + 6, 100), f"{round(val)}", ha="center", va="center", fontsize=9)

    ax.set_title(title, fontsize=12, pad=18)
    fig.tight_layout()
    return fig

def render_radar_png(pillar_scores: List[Tuple[str, float, Dict[str, int]]]) -> Optional[bytes]:
    if not HAS_MPL:
        return None
    try:
        fig = radar_chart(pillar_scores, title="Channel Readiness by Pillar")
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=180)
        plt.close(fig)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None

# =====
# DOCX generation - EXECUTIVE STYLE
# =====
def build_docx(
    company: str,
    name: str,
    email: str,
    role: str,
    phone: str,
    pillar_scores: List[Tuple[str, float, Dict[str, int]]],
    overall: float,
    brand_name: str,
    tsd_name: Optional[str],
    include_radar: bool = True,
    include_table: bool = True
) -> bytes:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    
    doc = Document()
    
    # Set narrow margins for more space
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # === PAGE 1: EXECUTIVE SUMMARY ===
    
    # Header with brand and date
    header = doc.add_paragraph()
    header_run = header.add_run(brand_name)
    header_run.font.size = Pt(18)
    header_run.font.name = "Calibri"
    header_run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    header_run.bold = True
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime('%B %d, %Y'))
    date_run.font.size = Pt(10)
    date_run.font.name = "Calibri"
    date_run.font.color.rgb = RGBColor(89, 89, 89)
    
    # Horizontal line
    doc.add_paragraph("_" * 80)
    
    # Company name - prominent
    company_para = doc.add_paragraph()
    company_run = company_para.add_run(f"{company}")
    company_run.font.size = Pt(22)
    company_run.font.name = "Calibri"
    company_run.bold = True
    company_run.font.color.rgb = RGBColor(0, 0, 0)
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Channel Readiness Assessment")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)
    
    if tsd_name:
        tsd_para = doc.add_paragraph()
        tsd_run = tsd_para.add_run(f"Technology Service Distributor: {tsd_name}")
        tsd_run.font.size = Pt(10)
        tsd_run.font.italic = True
    
    doc.add_paragraph("")
    
    # Executive Summary Box
    exec_summary = doc.add_paragraph()
    exec_summary_title = exec_summary.add_run("EXECUTIVE SUMMARY")
    exec_summary_title.font.size = Pt(12)
    exec_summary_title.font.name = "Calibri"
    exec_summary_title.bold = True
    exec_summary_title.font.color.rgb = RGBColor(0, 51, 102)
    
    # Score - Large and prominent
    _tier = tier_for(overall)
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    score_run = score_para.add_run(f"{round(overall)}")
    score_run.font.size = Pt(48)
    score_run.font.name = "Calibri"
    score_run.bold = True
    score_run.font.color.rgb = RGBColor(0, 102, 204)
    
    score_label = doc.add_paragraph()
    score_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = score_label.add_run(f"Channel Readiness Score\n{_tier} Maturity")
    label_run.font.size = Pt(14)
    label_run.font.name = "Calibri"
    label_run.font.color.rgb = RGBColor(89, 89, 89)
    
    doc.add_paragraph("")
    
    # Contact info - clean table
    contact_table = doc.add_table(rows=4, cols=2)
    contact_table.style = 'Light Grid Accent 1'
    
    contact_data = [
        ("Contact", name),
        ("Email", email),
        ("Title", role),
        ("Phone", phone if phone else "—")
    ]
    
    for i, (label, value) in enumerate(contact_data):
        contact_table.cell(i, 0).text = label
        contact_table.cell(i, 1).text = value
        contact_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        contact_table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(10)
        contact_table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph("")
    
    # Radar chart - larger and prominent
    if include_radar:
        png_bytes = render_radar_png(pillar_scores)
        if png_bytes:
            radar_hdr = doc.add_paragraph()
            radar_title = radar_hdr.add_run("Capability Assessment Radar")
            radar_title.font.size = Pt(12)
            radar_title.bold = True
            radar_title.font.color.rgb = RGBColor(0, 51, 102)
            
            radar_para = doc.add_paragraph()
            radar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
    
    # Page break
    doc.add_page_break()
    
    # === PAGE 2: DETAILED ANALYSIS ===
    
    # Page 2 header
    page2_header = doc.add_paragraph()
    p2h_run = page2_header.add_run("Detailed Assessment Results")
    p2h_run.font.size = Pt(16)
    p2h_run.bold = True
    p2h_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph("_" * 80)
    doc.add_paragraph("")
    
    # Pillar Scores - Professional table
    pillar_hdr = doc.add_paragraph()
    pillar_title = pillar_hdr.add_run("Pillar Performance Summary")
    pillar_title.font.size = Pt(12)
    pillar_title.bold = True
    pillar_title.font.color.rgb = RGBColor(0, 51, 102)
    
    # Create professional table
    score_table = doc.add_table(rows=len(pillar_scores) + 1, cols=3)
    score_table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ["Pillar", "Score", "Assessment"]
    for i, header_text in enumerate(headers):
        cell = score_table.cell(0, i)
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        # Add shading to header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9E2F3')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for i, (pname, pscore, _) in enumerate(pillar_scores, start=1):
        score_table.cell(i, 0).text = pname
        score_table.cell(i, 1).text = str(round(pscore))
        
        # Color code the score
        score_cell = score_table.cell(i, 1).paragraphs[0].runs[0]
        if pscore >= 80:
            score_cell.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif pscore >= 60:
            score_cell.font.color.rgb = RGBColor(0, 102, 204)  # Blue
        elif pscore >= 40:
            score_cell.font.color.rgb = RGBColor(255, 140, 0)  # Orange
        else:
            score_cell.font.color.rgb = RGBColor(204, 0, 0)  # Red
        score_cell.bold = True
        
        # Assessment text
        if pscore >= 80:
            assessment = "Strong"
        elif pscore >= 60:
            assessment = "Solid Foundation"
        elif pscore >= 40:
            assessment = "Emerging"
        else:
            assessment = "Needs Development"
        score_table.cell(i, 2).text = assessment
        score_table.cell(i, 2).paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph("")
    
    # Key Findings
    findings_hdr = doc.add_paragraph()
    findings_title = findings_hdr.add_run("Key Findings")
    findings_title.font.size = Pt(12)
    findings_title.bold = True
    findings_title.font.color.rgb = RGBColor(0, 51, 102)
    
    strengths, gaps = derive_strengths_gaps(pillar_scores)
    
    # Strengths
    strength_para = doc.add_paragraph()
    strength_label = strength_para.add_run("Areas of Strength:")
    strength_label.font.bold = True
    strength_label.font.size = Pt(11)
    strength_label.font.color.rgb = RGBColor(0, 128, 0)
    
    for s in strengths:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.text = s
        bullet.runs[0].font.size = Pt(10)
    
    doc.add_paragraph("")
    
    # Development areas
    gap_para = doc.add_paragraph()
    gap_label = gap_para.add_run("Development Priorities:")
    gap_label.font.bold = True
    gap_label.font.size = Pt(11)
    gap_label.font.color.rgb = RGBColor(204, 102, 0)
    
    for g in gaps:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.text = g
        bullet.runs[0].font.size = Pt(10)
    
    doc.add_paragraph("")
    
    # Recommendations - Numbered and executive style
    rec_hdr = doc.add_paragraph()
    rec_title = rec_hdr.add_run("Strategic Recommendations (Next 90 Days)")
    rec_title.font.size = Pt(12)
    rec_title.bold = True
    rec_title.font.color.rgb = RGBColor(0, 51, 102)
    
    recs = recommend_actions(pillar_scores)
    for idx, rec in enumerate(recs, 1):
        rec_para = doc.add_paragraph(style='List Number')
        rec_para.text = rec
        rec_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Call to action - boxed
    cta_para = doc.add_paragraph()
    cta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cta_run = cta_para.add_run("Ready to achieve a 90+ Channel Readiness Score?")
    cta_run.font.size = Pt(12)
    cta_run.bold = True
    cta_run.font.color.rgb = RGBColor(0, 51, 102)
    
    cta2 = doc.add_paragraph()
    cta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cta2_run = cta2.add_run("Schedule a comprehensive XplainIQ GTM Assessment")
    cta2_run.font.size = Pt(11)
    cta2_run.font.color.rgb = RGBColor(89, 89, 89)
    
    # Footer
    doc.add_paragraph("")
    footer_line = doc.add_paragraph("_" * 80)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("© Innovative Networx – XplainIQ™ | Confidential & Proprietary")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# =====
# Email generation functions
# =====
def generate_html_email(
    company: str,
    name: str,
    pillar_scores: List[Tuple[str, float, Dict[str, int]]],
    overall: float,
    tier: str
) -> str:
    """Generate beautiful HTML email for client follow-up"""
    
    strengths, gaps = derive_strengths_gaps(pillar_scores)
    
    # Generate radar chart as base64 for embedding
    radar_base64 = ""
    if HAS_MPL:
        try:
            png_bytes = render_radar_png(pillar_scores)
            if png_bytes:
                radar_base64 = base64.b64encode(png_bytes).decode('utf-8')
        except Exception:
            pass
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: 'Segoe UI', Calibri, Arial, sans-serif; margin: 0; padding: 0; background-color: #f5f5f5; }}
        .container {{ max-width: 600px; margin: 40px auto; background: white; border-radius: 8px; overflow: hidden; box-shadow: 0 4px 20px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #003366 0%, #0066cc 100%); color: white; padding: 40px 30px; text-align: center; }}
        .header h1 {{ margin: 0; font-size: 28px; font-weight: bold; }}
        .header p {{ margin: 10px 0 0 0; font-size: 16px; opacity: 0.9; }}
        .content {{ padding: 40px 30px; }}
        .greeting {{ font-size: 16px; color: #333; margin-bottom: 20px; }}
        .score-box {{ background: linear-gradient(135deg, #e8f4fd 0%, #d0e8f7 100%); border-radius: 12px; padding: 30px; text-align: center; margin: 30px 0; border: 2px solid #0066cc; }}
        .score-number {{ font-size: 64px; font-weight: bold; color: #0066cc; margin: 0; line-height: 1; }}
        .score-label {{ font-size: 18px; color: #003366; margin-top: 10px; font-weight: 600; }}
        .section {{ margin: 30px 0; }}
        .section-title {{ font-size: 18px; font-weight: bold; color: #003366; margin-bottom: 15px; border-bottom: 2px solid #0066cc; padding-bottom: 8px; }}
        .radar-chart {{ text-align: center; margin: 25px 0; }}
        .radar-chart img {{ max-width: 100%; height: auto; border-radius: 8px; }}
        .insight-box {{ background: #f8f9fa; border-left: 4px solid #0066cc; padding: 15px 20px; margin: 15px 0; border-radius: 4px; }}
        .insight-box h4 {{ margin: 0 0 10px 0; color: #003366; font-size: 14px; }}
        .insight-list {{ margin: 0; padding-left: 20px; color: #555; line-height: 1.8; }}
        .cta-box {{ background: #003366; color: white; padding: 30px; text-align: center; margin: 30px 0; border-radius: 8px; }}
        .cta-box h3 {{ margin: 0 0 15px 0; font-size: 22px; }}
        .cta-button {{ display: inline-block; background: #0066cc; color: white; padding: 15px 40px; text-decoration: none; border-radius: 6px; font-weight: bold; font-size: 16px; margin-top: 10px; }}
        .cta-button:hover {{ background: #0052a3; }}
        .footer {{ background: #f8f9fa; padding: 25px 30px; text-align: center; color: #666; font-size: 13px; border-top: 1px solid #ddd; }}
        .footer-logo {{ font-weight: bold; color: #003366; font-size: 14px; margin-bottom: 10px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Your Channel Readiness Results</h1>
            <p>XplainIQ™ Assessment for {company}</p>
        </div>
        
        <div class="content">
            <div class="greeting">
                Hi {name},
            </div>
            
            <p style="color: #555; line-height: 1.6;">
                Thank you for completing the XplainIQ Channel Readiness Assessment. Our team has reviewed your responses, and we're pleased to share your results.
            </p>
            
            <div class="score-box">
                <div class="score-number">{round(overall)}</div>
                <div class="score-label">Channel Readiness Score</div>
                <div style="color: #666; margin-top: 8px; font-size: 16px;">{tier} Maturity Level</div>
            </div>
"""
    
    # Add radar chart if available
    if radar_base64:
        html += f"""
            <div class="section">
                <div class="section-title">Your Capability Profile</div>
                <div class="radar-chart">
                    <img src="data:image/png;base64,{radar_base64}" alt="Channel Readiness Radar" />
                </div>
            </div>
"""
    
    # Add key insights
    html += f"""
            <div class="section">
                <div class="section-title">Key Insights</div>
                
                <div class="insight-box" style="border-left-color: #00aa00;">
                    <h4>✓ Areas of Strength</h4>
                    <ul class="insight-list">
"""
    
    for strength in strengths[:2]:
        html += f"                        <li>{strength}</li>\n"
    
    html += f"""
                    </ul>
                </div>
                
                <div class="insight-box" style="border-left-color: #ff8800;">
                    <h4>→ Development Opportunities</h4>
                    <ul class="insight-list">
"""
    
    for gap in gaps[:3]:
        html += f"                        <li>{gap}</li>\n"
    
    html += f"""
                    </ul>
                </div>
            </div>
            
            <p style="color: #555; line-height: 1.6; margin-top: 30px;">
                Your {tier.lower()} maturity level indicates a solid foundation with clear paths to optimization. Our advisors have identified specific strategies that can accelerate your channel growth.
            </p>
            
            <div class="cta-box">
                <h3>Ready to Reach 90+ and Scale Your Channel?</h3>
                <p style="margin: 15px 0; opacity: 0.95;">
                    Schedule a complimentary 30-minute strategy session with our GTM advisors to discuss your personalized roadmap.
                </p>
                <a href="mailto:info@innovativenetworx.com?subject=Channel Assessment Follow-up - {company}" class="cta-button">
                    Book Your Strategy Session
                </a>
            </div>
            
            <p style="color: #555; line-height: 1.6; margin-top: 30px; font-size: 14px;">
                Questions about your results? Simply reply to this email, and one of our channel strategy advisors will be in touch within 24 hours.
            </p>
            
            <p style="color: #555; margin-top: 30px;">
                Best regards,<br>
                <strong>The XplainIQ Team</strong><br>
                Innovative Networx
            </p>
        </div>
        
        <div class="footer">
            <div class="footer-logo">XplainIQ™ by Innovative Networx</div>
            <div>Engineering Predictable Go-To-Market Outcomes</div>
            <div style="margin-top: 10px;">
                © {datetime.now().year} Innovative Networx. Confidential & Proprietary.
            </div>
        </div>
    </div>
</body>
</html>"""
    
    return html


def generate_text_email(
    company: str,
    name: str,
    pillar_scores: List[Tuple[str, float, Dict[str, int]]],
    overall: float,
    tier: str
) -> str:
    """Generate plain text email for client follow-up"""
    
    strengths, gaps = derive_strengths_gaps(pillar_scores)
    recs = recommend_actions(pillar_scores)
    
    text = f"""Subject: Your Channel Readiness Assessment Results - {company}

Hi {name},

Thank you for completing the XplainIQ Channel Readiness Assessment. Our team has reviewed your responses, and I'm pleased to share your results.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

YOUR CHANNEL READINESS SCORE: {round(overall)} / 100
Maturity Level: {tier}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PILLAR SCORES:
"""
    
    for pname, pscore, _ in pillar_scores:
        text += f"  • {pname}: {round(pscore)}\n"
    
    text += f"""
KEY INSIGHTS:

✓ Areas of Strength:
"""
    
    for strength in strengths:
        text += f"  • {strength}\n"
    
    text += f"""
→ Development Opportunities:
"""
    
    for gap in gaps:
        text += f"  • {gap}\n"
    
    text += f"""
TOP 3 STRATEGIC RECOMMENDATIONS (Next 90 Days):

"""
    
    for i, rec in enumerate(recs, 1):
        text += f"{i}. {rec}\n\n"
    
    text += f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

NEXT STEPS:

Your {tier.lower()} maturity level indicates a solid foundation with clear paths to optimization. Our advisors have identified specific strategies that can accelerate your channel growth.

Ready to reach 90+ and scale your channel?

Schedule a complimentary 30-minute strategy session with our GTM advisors to discuss your personalized roadmap.

Reply to this email or contact us at:
info@innovativenetworx.com

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Questions about your results? Simply reply to this email, and one of our channel strategy advisors will be in touch within 24 hours.

Best regards,
The XplainIQ Team
Innovative Networx

XplainIQ™ | Engineering Predictable Go-To-Market Outcomes
© {datetime.now().year} Innovative Networx. Confidential & Proprietary.
"""
    
    return text


def generate_pdf_report(
    company: str,
    name: str,
    pillar_scores: List[Tuple[str, float, Dict[str, int]]],
    overall: float,
    tier: str
) -> Optional[bytes]:
    """Generate PDF from HTML email template using weasyprint or reportlab fallback"""
    
    try:
        # Try using weasyprint (best quality, but requires system libraries)
        from weasyprint import HTML
        
        html_content = generate_html_email(company, name, pillar_scores, overall, tier)
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes
        
    except ImportError:
        # Fallback: Try reportlab (pure Python, always works)
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#003366'),
                spaceAfter=6,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Normal'],
                fontSize=14,
                textColor=colors.HexColor('#666666'),
                spaceAfter=20,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#003366'),
                spaceAfter=12,
                spaceBefore=20,
                fontName='Helvetica-Bold'
            )
            
            # Title
            story.append(Paragraph("Channel Readiness Assessment", title_style))
            story.append(Paragraph(f"{company}", subtitle_style))
            
            # TSD name if provided
            tsd_name_str = ""
            try:
                row_data = st.session_state.get('current_selected_row', {})
                tsd_val = row_data.get('tsd_request_name', '')
                if pd.notna(tsd_val) and tsd_val:
                    tsd_name_str = str(tsd_val)
            except:
                pass
            
            if tsd_name_str:
                tsd_style = ParagraphStyle(
                    'TSD',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=colors.HexColor('#666666'),
                    spaceAfter=20,
                    alignment=TA_CENTER,
                    fontName='Helvetica-Oblique'
                )
                story.append(Paragraph(f"Technology Service Distributor: {tsd_name_str}", tsd_style))
            else:
                story.append(Spacer(1, 0.2*inch))
            
            # Score box
            score_data = [
                ['Channel Readiness Score'],
                [f'{round(overall)}'],
                [f'{tier} Maturity']
            ]
            score_table = Table(score_data, colWidths=[4*inch])
            score_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, 1), 48),
                ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
                ('TEXTCOLOR', (0, 1), (-1, 1), colors.HexColor('#0066cc')),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
                ('TOPPADDING', (0, 0), (-1, -1), 15),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#0066cc')),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#e8f4fd')),
            ]))
            story.append(score_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Contact info
            story.append(Paragraph("Contact Information", heading_style))
            contact_data = [
                ['Contact:', name],
                ['Company:', company]
            ]
            contact_table = Table(contact_data, colWidths=[1.5*inch, 4*inch])
            contact_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(contact_table)
            story.append(Spacer(1, 0.2*inch))
            
            # Pillar scores
            story.append(Paragraph("Pillar Performance", heading_style))
            pillar_data = [['Pillar', 'Score']]
            for pname, pscore, _ in pillar_scores:
                pillar_data.append([pname, str(round(pscore))])
            
            pillar_table = Table(pillar_data, colWidths=[4*inch, 1*inch])
            pillar_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#d9e2f3')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ]))
            story.append(pillar_table)
            story.append(Spacer(1, 0.2*inch))
            
            # Key insights
            strengths, gaps = derive_strengths_gaps(pillar_scores)
            
            story.append(Paragraph("Key Insights", heading_style))
            story.append(Paragraph("<b>Areas of Strength:</b>", styles['Normal']))
            for strength in strengths:
                story.append(Paragraph(f"• {strength}", styles['Normal']))
            
            story.append(Spacer(1, 0.15*inch))
            story.append(Paragraph("<b>Development Opportunities:</b>", styles['Normal']))
            for gap in gaps:
                story.append(Paragraph(f"• {gap}", styles['Normal']))
            
            story.append(Spacer(1, 0.3*inch))
            
            # Strategic Recommendations
            story.append(Paragraph("Strategic Recommendations (Next 90 Days)", heading_style))
            recs = recommend_actions(pillar_scores)
            for i, rec in enumerate(recs, 1):
                rec_para = Paragraph(f"{i}. {rec}", styles['Normal'])
                story.append(rec_para)
                story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.3*inch))
            
            # CTA
            cta_style = ParagraphStyle(
                'CTA',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#003366'),
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                spaceAfter=8
            )
            story.append(Paragraph("Ready to achieve a 90+ Channel Readiness Score?", cta_style))
            
            cta_sub_style = ParagraphStyle(
                'CTASub',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#666666'),
                alignment=TA_CENTER,
                spaceAfter=20
            )
            story.append(Paragraph("Schedule a comprehensive XplainIQ GTM Assessment", cta_sub_style))
            
            # Footer
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#666666'),
                alignment=TA_CENTER
            )
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph("_" * 80, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph("© Innovative Networx – XplainIQ™ | Confidential & Proprietary", footer_style))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.read()
            
        except ImportError:
            # Neither library available
            return None
    
    except Exception:
        return None
def init_session_storage():
    if 'leads_db' not in st.session_state:
        st.session_state.leads_db = []
        try:
            csv_path = APP_DIR / "leads.csv"
            if csv_path.exists():
                df = pd.read_csv(csv_path)
                st.session_state.leads_db = df.to_dict('records')
        except Exception:
            pass

def persist_lead(row: Dict) -> Tuple[bool, str]:
    try:
        if 'leads_db' not in st.session_state:
            st.session_state.leads_db = []
        st.session_state.leads_db.append(row)
        
        try:
            path = APP_DIR / "leads.csv"
            df = pd.DataFrame(st.session_state.leads_db)
            df.to_csv(path, index=False)
            return True, f"Saved to session + CSV backup at {path.resolve()}"
        except Exception as e:
            return True, f"Saved to session (CSV backup failed: {e})"
    except Exception as e:
        return False, f"Error saving: {e}"

def get_leads_dataframe() -> Optional[pd.DataFrame]:
    if 'leads_db' in st.session_state and st.session_state.leads_db:
        try:
            df = pd.DataFrame(st.session_state.leads_db)
            if 'ts' in df.columns:
                df['ts'] = pd.to_datetime(df['ts'], errors='coerce')
                df = df.sort_values('ts', ascending=False)
            return df
        except Exception:
            pass
    return None

def update_lead_status(index: int, new_status: str):
    """Update the status of a lead"""
    if 'leads_db' in st.session_state and index < len(st.session_state.leads_db):
        st.session_state.leads_db[index]['status'] = new_status
        try:
            path = APP_DIR / "leads.csv"
            df = pd.DataFrame(st.session_state.leads_db)
            df.to_csv(path, index=False)
        except Exception:
            pass

# =====
# Query params
# =====
def _get_query_params() -> dict[str, list[str]]:
    try:
        qp = st.query_params
        if isinstance(qp, dict):
            return {k: ([v] if isinstance(v, str) else v) for k, v in qp.items()}
        return {k: ([v] if isinstance(v, str) else v) for k, v in qp.items()}
    except Exception:
        return {}

def get_query_param(key: str, default: str = "") -> str:
    qp = _get_query_params()
    vals = qp.get(key, [default]) or [default]
    return vals[0]

def prefill_answers_from_query() -> Dict[str, int]:
    out = {}
    for qid in QUESTIONS.keys():
        try:
            val = int(get_query_param(qid.lower(), ""))
            if 1 <= val <= 5:
                out[qid] = val
        except Exception:
            pass
    return out

# =====
# Initialize
# =====
init_session_storage()

# =====
# Determine mode
# =====
admin_flag = str(get_query_param("admin", "0")).lower() in ("1", "true", "yes")
debug_mode = str(get_query_param("debug", "0")).lower() in ("1", "true", "yes")

# =====
# Header
# =====
show_logo_any(ASSET_LOGO_PATH, width=220, show_debug=debug_mode)
st.title("XplainIQ lite: Channel Readiness Scoring Index")

# URL Prefills
prefill_company = get_query_param("company", "")
prefill_name    = get_query_param("name", "")
prefill_email   = get_query_param("email", "")
prefill_role    = get_query_param("role", "")
prefill_phone   = get_query_param("phone", "")
prefill_tsd     = get_query_param("tsd", "")
prefilled_qs    = prefill_answers_from_query()

if admin_flag:
    st.caption("🔓 Admin mode active")

# =====
# CLIENT MODE
# =====
if not admin_flag:
    st.write("Please answer the questions below. Our advisors will review your submission and send your personalized report via email within 2-3 business days.")

    with st.sidebar:
        show_logo_any(ASSET_LOGO_PATH, width=180, show_debug=debug_mode)
        st.markdown(f"**{BRAND_NAME}**")
        st.markdown("---")
        st.markdown("### Contact Information")
        name  = st.text_input("Your Name*", value=prefill_name)
        email = st.text_input("Work Email*", value=prefill_email)
        role  = st.text_input("Title / Role*", value=prefill_role)
        phone = st.text_input("Phone (optional)", value=prefill_phone)
        tsd_name = st.text_input("TSD request name", value=prefill_tsd)
        consent = st.checkbox(
            "I consent to Innovative Networx reviewing my responses and contacting me about my results.",
            value=False
        )
        st.caption("* Required fields")
        if "last_submit_ts" not in st.session_state:
            st.session_state.last_submit_ts = 0.0

    st.markdown("#### Company")
    company = st.text_input("Company Name*", value=prefill_company)

    st.markdown("### Assessment Questions")
    st.caption("Rate each statement from 1 (Strongly Disagree) to 5 (Strongly Agree)")
    
    answers: Dict[str, int] = {**prefilled_qs}
    for qid, text in QUESTIONS.items():
        default_val = int(answers.get(qid, 3))
        answers[qid] = st.slider(f"**{qid}:** {text}", min_value=1, max_value=5, value=default_val, key=f"q_{qid}")

    st.markdown("---")
    
    if st.button("📤 Submit Assessment", type="primary", use_container_width=True):
        now = time.time()
        if now - st.session_state.last_submit_ts < 60:
            st.warning("⏳ Please wait a minute before submitting again.")
        elif not consent:
            st.error("❌ Please provide consent to proceed.")
        elif not email or "@" not in email:
            st.error("❌ Please enter a valid work email.")
        elif not name or not role or not company:
            st.error("❌ Please fill in all required fields (Name, Email, Role, Company).")
        else:
            st.session_state.last_submit_ts = now
            
            # Compute scores
            pillar_scores, overall = compute_scores(answers)
            tier = tier_for(overall)
            
            payload = {
                "ts": datetime.utcnow().isoformat() + "Z",
                "brand_name": BRAND_NAME,
                "tsd_request_name": tsd_name if tsd_name else "",
                "company": company,
                "name": name,
                "email": email,
                "role": role,
                "phone": phone if phone else "",
                "score_overall": round(overall),
                "tier": tier,
                "pillar_scores": str({p[0]: round(p[1]) for p in pillar_scores}),
                "answers": str(answers),
                "status": "Pending Review",
            }
            
            ok, msg = persist_lead(payload)
            if ok:
                st.success("✅ **Thank you for your submission!**")
                st.info("📧 Our advisors will review your responses and send your personalized Channel Readiness Report via email within 2-3 business days.")
                st.balloons()
            else:
                st.error(f"❌ Submission error: {msg}")

    st.caption("Powered by XplainIQ™ • Engineering Predictable Go-To-Market Outcomes.")

# =====
# ADMIN MODE
# =====
else:
    st.write("Admin panel for reviewing submissions and generating reports.")
    
    st.markdown("---")
    st.subheader("📊 Latest Submissions")
    
    col_refresh, col_space = st.columns([1, 5])
    with col_refresh:
        if st.button("🔄 Refresh"):
            st.rerun()
    
    df = get_leads_dataframe()
    if df is not None and not df.empty:
        # Display table
        display_cols = [c for c in ["ts","company","name","email","score_overall","tier","status"] if c in df.columns]
        display_cols += [c for c in df.columns if c not in display_cols]
        st.dataframe(df[display_cols].head(25), use_container_width=True)
        
        # Export CSV
        csv_export = df.to_csv(index=False)
        st.download_button(
            "📥 Export All Leads (CSV)",
            data=csv_export,
            file_name=f"leads_export_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv"
        )
        
        st.markdown("---")
        st.subheader("📄 Generate Report from Submission")
        
        # Create dropdown
        submission_labels = []
        for idx, row in df.iterrows():
            label = f"{row['company']} - {row['name']} ({row['email']}) - {row['ts']}"
            submission_labels.append(label)
        
        selected_idx = st.selectbox(
            "Select a submission:",
            range(len(df)),
            format_func=lambda i: submission_labels[i]
        )
        
        # Show selected submission details
        selected_row = df.iloc[selected_idx]
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Score", f"{selected_row['score_overall']}/100")
        with col2:
            st.metric("Tier", selected_row['tier'])
        with col3:
            current_status = selected_row.get('status', 'Pending Review')
            st.metric("Status", current_status)
        
        # Status update
        new_status = st.selectbox(
            "Update status:",
            ["Pending Review", "Report Generated", "Report Sent"],
            index=["Pending Review", "Report Generated", "Report Sent"].index(current_status) if current_status in ["Pending Review", "Report Generated", "Report Sent"] else 0,
            key="status_select"
        )
        
        if st.button("💾 Update Status"):
            update_lead_status(selected_idx, new_status)
            st.success(f"✅ Status updated to: {new_status}")
            st.rerun()
        
        # Parse data once for all report types
        try:
            import ast
            stored_answers = ast.literal_eval(selected_row['answers']) if isinstance(selected_row['answers'], str) else selected_row['answers']
            stored_pillar_scores_dict = ast.literal_eval(selected_row['pillar_scores']) if isinstance(selected_row['pillar_scores'], str) else selected_row['pillar_scores']
            
            # Reconstruct pillar_scores
            reconstructed_pillar_scores = []
            for pname, qids in PILLARS:
                pscore = stored_pillar_scores_dict.get(pname, 0)
                detail = {q: stored_answers.get(q, 0) for q in qids}
                reconstructed_pillar_scores.append((pname, pscore, detail))
        except Exception as e:
            st.error(f"Error parsing data: {e}")
            reconstructed_pillar_scores = []
        
        # Four-column layout for report options
        st.markdown("---")
        st.subheader("📬 Delivery Options")
        
        col1, col2, col3, col4 = st.columns(4)
        
        # Option 1: DOCX Download
        with col1:
            if st.button("📄 Generate DOCX", use_container_width=True, type="primary"):
                try:
                    if HAS_DOCX:
                        report_docx = build_docx(
                            company=str(selected_row['company']),
                            name=str(selected_row['name']),
                            email=str(selected_row['email']),
                            role=str(selected_row['role']),
                            phone=str(selected_row.get('phone', '')),
                            pillar_scores=reconstructed_pillar_scores,
                            overall=float(selected_row['score_overall']),
                            brand_name=BRAND_NAME,
                            tsd_name=str(selected_row.get('tsd_request_name', '')) if pd.notna(selected_row.get('tsd_request_name')) else None,
                            include_radar=True,
                            include_table=True,
                        )
                        
                        st.session_state['generated_docx'] = report_docx
                        st.session_state['docx_filename'] = f"{str(selected_row['company']).replace(' ', '')}_ChannelReadiness_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                        st.success("✅ DOCX ready!")
                    else:
                        st.error("❌ DOCX library not available")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
        
        # Option 2: HTML Email
        with col2:
            if st.button("📧 Copy HTML Email", use_container_width=True):
                try:
                    html_email = generate_html_email(
                        company=str(selected_row['company']),
                        name=str(selected_row['name']),
                        pillar_scores=reconstructed_pillar_scores,
                        overall=float(selected_row['score_overall']),
                        tier=str(selected_row['tier'])
                    )
                    st.session_state['html_email'] = html_email
                    st.success("✅ HTML ready to copy!")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
        
        # Option 3: Plain Text Email
        with col3:
            if st.button("📋 Copy Text Email", use_container_width=True):
                try:
                    text_email = generate_text_email(
                        company=str(selected_row['company']),
                        name=str(selected_row['name']),
                        pillar_scores=reconstructed_pillar_scores,
                        overall=float(selected_row['score_overall']),
                        tier=str(selected_row['tier'])
                    )
                    st.session_state['text_email'] = text_email
                    st.success("✅ Text ready to copy!")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
        
        # Option 4: PDF Download
        with col4:
            if st.button("📑 Generate PDF", use_container_width=True):
                try:
                    # Store row data for PDF generation to access TSD name
                    st.session_state['current_selected_row'] = selected_row.to_dict()
                    
                    pdf_bytes = generate_pdf_report(
                        company=str(selected_row['company']),
                        name=str(selected_row['name']),
                        pillar_scores=reconstructed_pillar_scores,
                        overall=float(selected_row['score_overall']),
                        tier=str(selected_row['tier'])
                    )
                    if pdf_bytes:
                        st.session_state['generated_pdf'] = pdf_bytes
                        st.session_state['pdf_filename'] = f"{str(selected_row['company']).replace(' ', '')}_Assessment_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
                        st.success("✅ PDF ready!")
                    else:
                        st.error("❌ PDF generation requires additional libraries")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
        
        # Display outputs
        if 'generated_docx' in st.session_state:
            st.download_button(
                "⬇️ Download DOCX Report",
                data=st.session_state['generated_docx'],
                file_name=st.session_state['docx_filename'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx"
            )
        
        if 'generated_pdf' in st.session_state:
            st.download_button(
                "⬇️ Download PDF Report",
                data=st.session_state['generated_pdf'],
                file_name=st.session_state['pdf_filename'],
                mime="application/pdf",
                key="download_pdf"
            )
        
        if 'html_email' in st.session_state:
            with st.expander("📧 HTML Email (Copy & Paste into Email Client)", expanded=True):
                st.code(st.session_state['html_email'], language='html')
                st.caption("💡 Tip: Copy all text above and paste into your email client's HTML view")
        
        if 'text_email' in st.session_state:
            with st.expander("📋 Plain Text Email (Copy & Paste)", expanded=True):
                st.text_area("Email Content", st.session_state['text_email'], height=400, key="text_display")
                st.caption("💡 Tip: Copy the text above for LinkedIn, SMS, or plain text emails")
        
        # Preview scores
        with st.expander("🔍 Preview Scores & Recommendations"):
            try:
                import ast
                stored_answers = ast.literal_eval(selected_row['answers']) if isinstance(selected_row['answers'], str) else selected_row['answers']
                stored_pillar_scores_dict = ast.literal_eval(selected_row['pillar_scores']) if isinstance(selected_row['pillar_scores'], str) else selected_row['pillar_scores']
                
                reconstructed_pillar_scores = []
                for pname, qids in PILLARS:
                    pscore = stored_pillar_scores_dict.get(pname, 0)
                    detail = {q: stored_answers.get(q, 0) for q in qids}
                    reconstructed_pillar_scores.append((pname, pscore, detail))
                
                for pname, pscore, detail in reconstructed_pillar_scores:
                    st.write(f"**{pname}:** {round(pscore)}/100")
                    st.caption(pillar_commentary(pname, pscore))
                
                if HAS_MPL:
                    fig = radar_chart(reconstructed_pillar_scores, title="Channel Readiness by Pillar")
                    st.pyplot(fig, use_container_width=True)
                
            except Exception as e:
                st.error(f"Error displaying preview: {e}")
                
    else:
        st.info("📭 No submissions yet. Ask a client to submit their assessment.")
    
    st.caption("Powered by XplainIQ™ • Engineering Predictable Go-To-Market Outcomes.")
